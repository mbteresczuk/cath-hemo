"""
Populate Word document templates with cath report data.
"""
import io
import os
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches
from PIL import Image

BASE_DIR = Path(__file__).parent.parent


def _save_img_temp(img: Image.Image) -> str:
    """Save PIL image to a temporary PNG path."""
    import tempfile
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    img.save(tmp.name, format="PNG")
    return tmp.name


def populate_template(
    template_name: str,
    narrative: str,
    patient_data: dict,
    calculations: dict,
    annotated_image: Image.Image,
) -> bytes:
    """
    Fill in a Word template and return the resulting docx as bytes.

    template_name: 'standard' | 'pHTN' | 'OHT'
    """
    template_map = {
        "standard": "Report_template.docx",
        "pHTN": "Procedure Template - pHTN Drug Study.docx",
        "OHT": "Report Template - OHT.docx",
    }
    filename = template_map.get(template_name, "Report_template.docx")
    template_path = BASE_DIR / "templates" / filename

    if not template_path.exists():
        # Fall back to standard
        template_path = BASE_DIR / "templates" / "Report_template.docx"

    doc = Document(str(template_path))

    # Build replacement map for placeholder tokens
    today = datetime.today().strftime("%m/%d/%Y")
    patient_name = patient_data.get("name", "")
    mrn = patient_data.get("mrn", "")

    replacements = {
        "***": "___",  # generic placeholder
        "888%": "___",
    }

    # Apply replacements in all paragraphs and table cells
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            for run in para.runs:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)

    # Add annotated diagram and narrative at the end
    doc.add_page_break()
    doc.add_heading("Hemodynamic Diagram", level=2)

    img_path = _save_img_temp(annotated_image)
    try:
        doc.add_picture(img_path, width=Inches(5.5))
    except Exception:
        pass
    finally:
        os.unlink(img_path)

    doc.add_heading("Hemodynamic Findings", level=2)
    # Narrative is 4 paragraphs separated by \n\n — add each as a distinct
    # Word paragraph so they render with proper spacing in the docx.
    for para_text in narrative.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            doc.add_paragraph(para_text)

    # Add calculated values table
    doc.add_heading("Calculated Hemodynamics", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Value"
    hdr[2].text = "Units"

    calc_rows = [
        ("Qs (Cardiac Index)", calculations.get("qs"), "L/min/m²"),
        ("Qp", calculations.get("qp"), "L/min/m²"),
        ("Qp:Qs", calculations.get("qp_qs"), ":1"),
        ("PVRi", calculations.get("pvri"), "iWU"),
        ("SVRi", calculations.get("svri"), "iWU"),
        ("Rp/Rs", calculations.get("rp_rs"), ""),
        ("Mixed Venous Sat", calculations.get("mixed_venous_sat"), "%"),
        ("Mean PCWP", calculations.get("mean_pcwp"), "mmHg"),
        ("TPG", calculations.get("tpg"), "mmHg"),
    ]

    for param, value, units in calc_rows:
        if value is not None:
            row = table.add_row().cells
            row[0].text = param
            row[1].text = str(value)
            row[2].text = units

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
